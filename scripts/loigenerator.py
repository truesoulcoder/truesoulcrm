
import pandas as pd
from docx import Document
from docx.shared import Pt
from datetime import datetime
import os

# CONFIGURATION
INPUT_CSV = 'leads.csv'  # Your input CSV file
OUTPUT_DIR = 'DOCX_output'  # Folder to save LOIs
DOM_THRESHOLD = 90

# Utility functions
def format_currency(value):
    try:
        return "${:,.2f}".format(float(value))
    except:
        return "$0.00"

def generate_docx_loi(data, output_path):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    doc.add_paragraph(f"{data['Address']}, {data['City']}, {data['State']} {data['ZipCode']}")
    doc.add_paragraph(f"{datetime.now().strftime('%m/%d/%Y')}")
    doc.add_paragraph(f"Dear {data['FullName']},")

    doc.add_paragraph(
        f"I am writing to express my interest in structuring an all-cash offer on the property located at "
        f"{data['Address']}, {data['City']}, {data['State']} {data['ZipCode']}.")

    doc.add_paragraph(
        "Based on market conditions, comparable sales, and property profile, I would like to propose the following terms:")

    doc.add_paragraph("Offer Summary:")
    doc.add_paragraph(f"     - Price: {data['WholesaleValue']}")
    doc.add_paragraph("     - Option Period: 7 days (excluding weekends and federal holidays)")
    doc.add_paragraph(f"     - Earnest Money Deposit (EMD): {data['EMD']}")
    doc.add_paragraph("     - Buyer’s Assignment Consideration (BAC): $10")
    doc.add_paragraph(f"     - Closing Date: On or before {data['ClosingDate']}")


    doc.add_paragraph("Offer Highlights:")
    doc.add_paragraph("     - As-Is Condition")
    doc.add_paragraph("     - Buyer Pays All Closing Costs")
    doc.add_paragraph("     - Quick Close Available")

    doc.add_paragraph("Title Company: Kristin Blay at Ghrist Law – Patten Title")

    doc.add_paragraph("I am only able to acquire a limited number of properties at a time. As such, offer is only valid for 48 hours after it is received.")
    
    doc.add_paragraph("Warm regards,")
    doc.add_paragraph("Chris Phillips")
    doc.add_paragraph("True Soul Partners LLC")
    doc.add_paragraph("📞 817.500.1440")

    doc.add_paragraph(
        "This Letter of Intent to Purchase Real Estate outlines general intentions and is not legally binding. "
        "Terms are subject to further negotiation and approval. No party is obligated until a formal agreement is executed.")

    doc.save(output_path)

# Main execution
df = pd.read_csv(INPUT_CSV, low_memory=False)
df['FullName'] = df['FirstName'].fillna('') + ' ' + df['LastName'].fillna('')
df['DOM'] = pd.to_numeric(df['MLS_Curr_DaysOnMarket'], errors='coerce')

# Filter only by DOM
df = df[(df['DOM'] >= DOM_THRESHOLD) & (df['Contact1Email_1'].astype(str).str.contains('@', na=False))]

os.makedirs(OUTPUT_DIR, exist_ok=True)

def parse_currency(value):
    try:
        return float(str(value).replace("$", "").replace(",", "").strip())
    except:
        return 0.0

from datetime import datetime, timedelta  # 🔥 add timedelta here

# inside your loop
for _, row in df.iterrows():
    filename = f"{row['PropertyAddress']} - LETTER OF INTENT.docx"
    filename = ''.join(c for c in filename if c.isalnum() or c in (' ', '-', '_', '.')).rstrip()
    output_path = os.path.join(OUTPUT_DIR, filename)
    wholesale_val = 0.95 * parse_currency(row['WholesaleValue'])
    closing_date = (datetime.now() + timedelta(days=14)).strftime('%m/%d/%Y')  # 💥 14-day closing logic

    generate_docx_loi({
        'FullName': row['FullName'],
        'Address': row['PropertyAddress'],
        'City': row['PropertyCity'],
        'State': row['PropertyState'],
        'ZipCode': row['PropertyPostalCode'],
        'WholesaleValue': format_currency(wholesale_val),
        'EMD': format_currency(wholesale_val * 0.01),
        'ClosingDate': closing_date  # 💼 include dynamic closing date
    }, output_path)


print(f"Generated {len(df)} LOIs based on DOM ≥ {DOM_THRESHOLD} in '{OUTPUT_DIR}' folder.")
